import base64
import csv
import io
import os

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


def is_image_file(file_name: str) -> bool:
    """判断文件是否为支持的图片格式。"""
    ext = os.path.splitext(file_name)[1].lower()
    return ext in IMAGE_EXTENSIONS


def load_image_as_base64(file_name: str, file_bytes: bytes) -> str:
    """将图片字节流转为 data:image/{ext};base64,{data} 格式。"""
    ext = os.path.splitext(file_name)[1].lower().lstrip(".")
    if ext == "jpg":
        ext = "jpeg"
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    return f"data:image/{ext};base64,{b64}"


def load_document(file_path: str) -> str:
    """根据文件后缀自动选择加载器，返回文本内容。"""
    ext = os.path.splitext(file_path)[1].lower()
    loaders = {
        ".txt": _load_txt,
        ".pdf": _load_pdf,
        ".csv": _load_csv,
        ".docx": _load_docx,
    }
    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {', '.join(loaders)}")
    return loader(file_path)


def load_document_from_bytes(file_name: str, file_bytes: bytes) -> str:
    """从上传的文件字节流加载文档。"""
    ext = os.path.splitext(file_name)[1].lower()
    loaders = {
        ".txt": _load_txt_bytes,
        ".pdf": _load_pdf_bytes,
        ".csv": _load_csv_bytes,
        ".docx": _load_docx_bytes,
    }
    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {', '.join(loaders)}")
    return loader(file_bytes)


# --- 文件路径加载 ---

def _load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _load_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _load_csv(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        return "\n".join(", ".join(row) for row in reader)


def _load_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)


# --- 字节流加载 ---

def _load_txt_bytes(data: bytes) -> str:
    return data.decode("utf-8")


def _load_pdf_bytes(data: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _load_csv_bytes(data: bytes) -> str:
    text = data.decode("utf-8")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(", ".join(row) for row in reader)


def _load_docx_bytes(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)
